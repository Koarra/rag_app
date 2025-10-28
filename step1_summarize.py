"""
STEP 1: Extract text from document and generate summary
Supports PDF (with OCR for scanned documents) and DOCX files

Usage: python step1_summarize.py <input_file.pdf>
Output: Creates summary.json with the document summary
"""

import sys
import json
from pathlib import Path
from pydantic import BaseModel, Field
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from llama_index.core.program import LLMTextCompletionProgram
from llama_index.llms.azure_openai import AzureOpenAI
from docx import Document
import PyPDF2

# OCR imports (optional - will work without OCR if not installed)
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Note: OCR not available. Install with: pip install pdf2image pytesseract")


def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def ocr_pdf_page(pdf_path, page_num):
    """Extract text from a PDF page using OCR"""
    if not OCR_AVAILABLE:
        return ""

    try:
        images = convert_from_path(
            pdf_path,
            first_page=page_num + 1,
            last_page=page_num + 1,
            dpi=300
        )

        if images:
            return pytesseract.image_to_string(images[0])
        return ""

    except Exception as e:
        print(f"OCR error on page {page_num + 1}: {e}")
        return ""


def ocr_entire_pdf(pdf_path):
    """Extract text from entire PDF using OCR"""
    if not OCR_AVAILABLE:
        return ""

    try:
        images = convert_from_path(pdf_path, dpi=300)
        text_parts = []

        for i, img in enumerate(images):
            print(f"OCR processing page {i + 1}/{len(images)}...")
            text = pytesseract.image_to_string(img)
            if text.strip():
                text_parts.append(text)

        return "\n\n".join(text_parts)

    except Exception as e:
        print(f"OCR error: {e}")
        return ""


def extract_text_from_pdf(file_path):
    """Extract text from PDF file (with OCR fallback for scanned PDFs)"""
    text_parts = []

    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)

        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()

            if text and text.strip():
                text_parts.append(text)
            elif OCR_AVAILABLE:
                print(f"Using OCR for page {page_num + 1}...")
                ocr_text = ocr_pdf_page(file_path, page_num)
                if ocr_text:
                    text_parts.append(ocr_text)

    # If no text extracted and OCR is available, try full OCR
    if not text_parts and OCR_AVAILABLE:
        print("No text found. Running full OCR on entire PDF...")
        text_parts.append(ocr_entire_pdf(file_path))

    return "\n\n".join(text_parts)


def extract_text(file_path):
    """Extract text from PDF or DOCX"""
    extension = Path(file_path).suffix.lower()

    if extension == '.docx':
        return extract_text_from_docx(file_path)
    elif extension == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        print(f"Unsupported file type: {extension}")
        return None


# Pydantic model
class DocumentSummary(BaseModel):
    summary: str = Field(description="Comprehensive summary of the document")


def summarize_document(text, llm):
    """Generate summary using LlamaIndex"""

    # Limit text to 15000 characters
    text_to_summarize = text[:15000]

    program = LLMTextCompletionProgram.from_defaults(
        output_cls=DocumentSummary,
        llm=llm,
        prompt_template_str="""You are an expert document analyst. Summarize documents clearly and accurately.

Provide a comprehensive summary of this document.
Include:
- Main purpose and context
- Key parties involved
- Important dates and amounts
- Critical actions or decisions

Document:
{document_text}
""",
        verbose=False
    )

    result = program(document_text=text_to_summarize)
    return result.summary


def create_combined_summary(output_folder, summary_files, llm):
    """Create a combined summary from multiple document summaries"""

    # Read all summaries
    summaries_data = []
    for summary_file in summary_files:
        with open(summary_file, "r", encoding="utf-8") as f:
            data = json.load(f)
            summaries_data.append({
                "file_name": data.get("file_name", "Unknown"),
                "summary": data.get("summary", "")
            })

    # Combine summaries into a single text
    combined_text = ""
    for i, data in enumerate(summaries_data, 1):
        combined_text += f"\n--- Document {i}: {Path(data['file_name']).name} ---\n"
        combined_text += data['summary']
        combined_text += "\n"

    # Limit combined text to 15000 characters
    combined_text = combined_text[:15000]

    program = LLMTextCompletionProgram.from_defaults(
        output_cls=DocumentSummary,
        llm=llm,
        prompt_template_str="""You are an expert document analyst. Analyze multiple document summaries and create a combined summary.

Identify and highlight:
- Common themes across all documents
- Shared entities (people, organizations)
- Related events or transactions
- Overlapping time periods
- Key differences between documents

Document Summaries:
{all_summaries}
""",
        verbose=False
    )

    result = program(all_summaries=combined_text)
    combined_summary = result.summary

    # Save combined summary
    combined_result = {
        "file_count": len(summaries_data),
        "files": [Path(data["file_name"]).name for data in summaries_data],
        "combined_summary": combined_summary
    }

    with open(output_folder / "combined_summary.json", "w", encoding="utf-8") as f:
        json.dump(combined_result, f, indent=2)

    print(f"Saved: {output_folder}/combined_summary.json")
    print(f"Combined {len(summaries_data)} document summaries")


def main():
    if len(sys.argv) < 2:
        print("Usage: python step1_summarize.py <input_file.pdf> [output_folder]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_folder = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(".")
    output_folder.mkdir(parents=True, exist_ok=True)

    print(f"\n=== STEP 1: SUMMARIZE DOCUMENT ===")
    print(f"Processing: {input_file}")
    print(f"Output folder: {output_folder}")

    # Extract text
    print("Extracting text...")
    text = extract_text(input_file)

    if not text:
        print("Failed to extract text")
        sys.exit(1)

    print(f"Extracted {len(text)} characters")

    # Save extracted text
    with open(output_folder / "extracted_text.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print(f"Saved: {output_folder}/extracted_text.txt")

    # Initialize Azure OpenAI LLM
    llm = AzureOpenAI(
        engine="gpt-4o-mini",
        use_azure_ad=True,
        azure_ad_token_provider=get_bearer_token_provider(
            DefaultAzureCredential(), "https://cognitiveservices.azure.com/.default"
        )
    )

    # Generate summary
    print("Generating summary...")
    summary = summarize_document(text, llm)

    # Save summary with filename
    input_filename = Path(input_file).stem  # Get filename without extension
    result = {
        "file_name": input_file,
        "summary": summary
    }

    summary_filename = f"summary_{input_filename}.json"
    with open(output_folder / summary_filename, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2)

    print(f"Saved: {output_folder}/{summary_filename}")
    print(f"\nSummary:\n{summary}")

    # Check if there are multiple summary files and create a combined summary
    summary_files = list(output_folder.glob("summary_*.json"))
    if len(summary_files) > 1:
        print(f"\nFound {len(summary_files)} summary files. Creating combined summary...")
        create_combined_summary(output_folder, summary_files, llm)

    print("\n=== STEP 1 COMPLETE ===\n")


if __name__ == "__main__":
    main()
