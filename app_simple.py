"""
Simple Document Processing Application
Upload PDF/DOCX documents and get:
1. Summary
2. List of persons and companies
3. Description of each entity
4. Risk flags for money laundering and sanctions evasion
"""

import streamlit as st
import os
import json
import tempfile
from pathlib import Path
from openai import OpenAI
from docx import Document
import PyPDF2


# =============================================================================
# DOCUMENT TEXT EXTRACTION
# =============================================================================

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    doc = Document(file_path)
    text_parts = []

    # Get text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Get text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    text_parts = []

    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)

        for page in pdf_reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)

    return "\n\n".join(text_parts)


def extract_text(file_path):
    """Extract text from PDF or DOCX file"""
    extension = Path(file_path).suffix.lower()

    if extension == '.docx':
        return extract_text_from_docx(file_path)
    elif extension == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        return None


# =============================================================================
# STAGE 1: SUMMARIZE DOCUMENT
# =============================================================================

def summarize_document(api_key, document_text):
    """Generate a summary of the document using OpenAI"""

    st.write("🔄 Generating summary...")

    client = OpenAI(api_key=api_key)

    # Limit text length
    text_to_summarize = document_text[:15000]

    prompt = f"""Provide a comprehensive summary of this document.
Include:
- Main purpose and context
- Key parties involved
- Important dates and amounts
- Critical actions or decisions
- Relevant details

Document:
{text_to_summarize}"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert document analyst. Summarize documents clearly and accurately."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=1000
    )

    summary = response.choices[0].message.content
    return summary


# =============================================================================
# STAGE 2: EXTRACT ENTITIES
# =============================================================================

def extract_entities(api_key, document_text):
    """Extract persons and companies from the document"""

    st.write("🔄 Extracting persons and companies...")

    client = OpenAI(api_key=api_key)

    # Limit text length
    text_to_analyze = document_text[:15000]

    prompt = f"""Extract all persons and companies from this document.

Guidelines:
- Persons: Full names of individuals (e.g., "John Smith", "Dr. Jane Doe")
- Companies: Company names, organizations, institutions (e.g., "ABC Corporation", "Ministry of Finance")
- Be thorough - extract all entities mentioned

Document:
{text_to_analyze}"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert at extracting entities from documents."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1,
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "entities",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "persons": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "companies": {
                            "type": "array",
                            "items": {"type": "string"}
                        }
                    },
                    "required": ["persons", "companies"],
                    "additionalProperties": False
                }
            }
        }
    )

    entities = json.loads(response.choices[0].message.content)
    return entities


# =============================================================================
# STAGE 3: DESCRIBE ENTITIES
# =============================================================================

def describe_entities(api_key, document_text, persons, companies):
    """Generate detailed descriptions for each entity"""

    st.write("🔄 Generating entity descriptions...")

    client = OpenAI(api_key=api_key)

    # Combine all entities (limit to 10 for simplicity)
    all_entities = persons[:5] + companies[:5]

    if not all_entities:
        return {}

    entity_names = ", ".join(all_entities)
    text_to_analyze = document_text[:12000]

    prompt = f"""Analyze these entities from the document: {entity_names}

For each entity provide:
1. Description based on the document
2. Their role or position
3. Key activities they're involved in
4. Related entities
5. Any financial details

Document:
{text_to_analyze}"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert analyst. Analyze entities based only on information in the document."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=2000,
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "entity_descriptions",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "entities": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "name": {"type": "string"},
                                    "type": {"type": "string"},
                                    "description": {"type": "string"},
                                    "role": {"type": "string"},
                                    "key_activities": {"type": "array", "items": {"type": "string"}},
                                    "related_entities": {"type": "array", "items": {"type": "string"}},
                                    "financial_details": {"type": "string"}
                                },
                                "required": ["name", "type", "description", "role", "key_activities", "related_entities", "financial_details"],
                                "additionalProperties": False
                            }
                        }
                    },
                    "required": ["entities"],
                    "additionalProperties": False
                }
            }
        }
    )

    result = json.loads(response.choices[0].message.content)

    # Convert list to dictionary for easier access
    descriptions = {}
    for entity in result.get('entities', []):
        descriptions[entity['name']] = entity

    return descriptions


# =============================================================================
# STAGE 4: ANALYZE RISKS
# =============================================================================

def analyze_document_risk(api_key, document_text):
    """Analyze the document for money laundering and sanctions evasion risks"""

    st.write("🔄 Analyzing document risks...")

    client = OpenAI(api_key=api_key)

    text_to_analyze = document_text[:15000]

    prompt = f"""Analyze this document for money laundering and sanctions evasion indicators.

MONEY LAUNDERING INDICATORS:
• Unusual transaction patterns
• Shell companies or front companies
• Layering activities
• Structuring/Smurfing
• High-risk jurisdictions
• Rapid movement of funds

SANCTIONS EVASION INDICATORS:
• Transactions with sanctioned countries
• Use of front companies
• Re-routing through third countries
• False documentation
• Sanctioned individuals/entities
• Prohibited goods

Document:
{text_to_analyze}"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a financial crime analyst. Analyze documents for money laundering and sanctions evasion."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1,
        max_tokens=1500,
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "document_risk",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "overall_risk_level": {"type": "string", "enum": ["high", "medium", "low", "none"]},
                        "money_laundering_risk": {"type": "string", "enum": ["high", "medium", "low", "none"]},
                        "sanctions_evasion_risk": {"type": "string", "enum": ["high", "medium", "low", "none"]},
                        "red_flags": {"type": "array", "items": {"type": "string"}},
                        "suspicious_patterns": {"type": "array", "items": {"type": "string"}},
                        "analysis_summary": {"type": "string"}
                    },
                    "required": ["overall_risk_level", "money_laundering_risk", "sanctions_evasion_risk", "red_flags", "suspicious_patterns", "analysis_summary"],
                    "additionalProperties": False
                }
            }
        }
    )

    document_risk = json.loads(response.choices[0].message.content)
    return document_risk


def analyze_entity_risks(api_key, entity_descriptions):
    """Analyze individual entities for criminal involvement"""

    st.write("🔄 Analyzing entity risks...")

    if not entity_descriptions:
        return []

    client = OpenAI(api_key=api_key)

    # Build entity context
    entity_contexts = []
    for name, info in list(entity_descriptions.items())[:10]:  # Limit to 10
        context = f"""Entity: {name}
Type: {info.get('type', 'unknown')}
Role: {info.get('role', 'unknown')}
Description: {info.get('description', 'No description')}
Activities: {', '.join(info.get('key_activities', []))}"""
        entity_contexts.append(context)

    all_contexts = "\n\n".join(entity_contexts)

    prompt = f"""Analyze these entities for involvement in money laundering or sanctions evasion.
Only flag entities with credible evidence.

{all_contexts}"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert in financial crime detection. Only flag entities with credible evidence."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1,
        max_tokens=2000,
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "entity_risks",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "flagged_entities": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "entity_name": {"type": "string"},
                                    "entity_type": {"type": "string"},
                                    "crimes_flagged": {"type": "array", "items": {"type": "string", "enum": ["money_laundering", "sanctions_evasion"]}},
                                    "risk_level": {"type": "string", "enum": ["high", "medium", "low"]},
                                    "confidence": {"type": "number"},
                                    "evidence": {"type": "array", "items": {"type": "string"}},
                                    "reasoning": {"type": "string"}
                                },
                                "required": ["entity_name", "entity_type", "crimes_flagged", "risk_level", "confidence", "evidence", "reasoning"],
                                "additionalProperties": False
                            }
                        }
                    },
                    "required": ["flagged_entities"],
                    "additionalProperties": False
                }
            }
        }
    )

    result = json.loads(response.choices[0].message.content)
    return result.get('flagged_entities', [])


# =============================================================================
# PROCESS ONE DOCUMENT
# =============================================================================

def process_document(api_key, file_path, file_name):
    """Process a single document through all 4 stages"""

    st.subheader(f"📄 Processing: {file_name}")

    # Extract text
    with st.spinner("Extracting text..."):
        document_text = extract_text(file_path)

        if not document_text:
            st.error("Failed to extract text from document")
            return None

        st.success(f"✓ Extracted {len(document_text)} characters")

    results = {
        "file_name": file_name,
        "text": document_text
    }

    # STAGE 1: Summary
    with st.expander("📝 Stage 1: Document Summary", expanded=True):
        summary = summarize_document(api_key, document_text)
        results["summary"] = summary
        st.write(summary)

    # STAGE 2: Extract entities
    with st.expander("👥 Stage 2: Extracted Entities", expanded=True):
        entities = extract_entities(api_key, document_text)
        results["entities"] = entities

        persons = entities.get('persons', [])
        companies = entities.get('companies', [])

        col1, col2 = st.columns(2)

        with col1:
            st.write(f"**Persons ({len(persons)}):**")
            for person in persons:
                st.write(f"- {person}")

        with col2:
            st.write(f"**Companies ({len(companies)}):**")
            for company in companies:
                st.write(f"- {company}")

    # STAGE 3: Describe entities
    with st.expander("📋 Stage 3: Entity Descriptions", expanded=True):
        descriptions = describe_entities(api_key, document_text, persons, companies)
        results["descriptions"] = descriptions

        if descriptions:
            for entity_name, info in descriptions.items():
                st.write(f"**{entity_name}** ({info['type']})")
                st.write(f"*Role:* {info['role']}")
                st.write(f"*Description:* {info['description']}")

                if info['key_activities']:
                    st.write(f"*Activities:* {', '.join(info['key_activities'])}")

                if info['financial_details']:
                    st.write(f"*Financial:* {info['financial_details']}")

                st.write("---")
        else:
            st.write("No entity descriptions generated")

    # STAGE 4: Risk analysis
    with st.expander("⚠️ Stage 4: Risk Assessment", expanded=True):

        # Document risk
        document_risk = analyze_document_risk(api_key, document_text)
        results["document_risk"] = document_risk

        st.write("**Document-Level Risk:**")

        # Risk indicators with colors
        risk_colors = {
            "high": "🔴",
            "medium": "🟡",
            "low": "🟢",
            "none": "⚪"
        }

        col1, col2, col3 = st.columns(3)

        with col1:
            overall = document_risk['overall_risk_level']
            st.metric("Overall Risk", f"{risk_colors.get(overall, '⚪')} {overall.upper()}")

        with col2:
            ml_risk = document_risk['money_laundering_risk']
            st.metric("Money Laundering", f"{risk_colors.get(ml_risk, '⚪')} {ml_risk.upper()}")

        with col3:
            se_risk = document_risk['sanctions_evasion_risk']
            st.metric("Sanctions Evasion", f"{risk_colors.get(se_risk, '⚪')} {se_risk.upper()}")

        st.write(f"**Analysis:** {document_risk['analysis_summary']}")

        if document_risk['red_flags']:
            st.write("**Red Flags:**")
            for flag in document_risk['red_flags']:
                st.write(f"- {flag}")

        if document_risk['suspicious_patterns']:
            st.write("**Suspicious Patterns:**")
            for pattern in document_risk['suspicious_patterns']:
                st.write(f"- {pattern}")

        # Entity risks
        flagged_entities = analyze_entity_risks(api_key, descriptions)
        results["flagged_entities"] = flagged_entities

        if flagged_entities:
            st.write("---")
            st.write(f"**Flagged Entities: {len(flagged_entities)}**")

            for entity in flagged_entities:
                st.write(f"### {entity['entity_name']} ({entity['entity_type']})")
                st.write(f"**Risk Level:** {entity['risk_level'].upper()} (Confidence: {entity['confidence']:.2f})")

                crimes = ', '.join(entity['crimes_flagged']).replace('_', ' ').title()
                st.write(f"**Crimes:** {crimes}")

                st.write(f"**Reasoning:** {entity['reasoning']}")

                st.write("**Evidence:**")
                for evidence in entity['evidence']:
                    st.write(f"- {evidence}")

                st.write("---")
        else:
            st.success("✓ No entities flagged for criminal activity")

    # Download button
    st.download_button(
        label="📥 Download Results (JSON)",
        data=json.dumps(results, indent=2),
        file_name=f"{Path(file_name).stem}_results.json",
        mime="application/json"
    )

    return results


# =============================================================================
# MAIN APPLICATION
# =============================================================================

def main():
    """Main Streamlit application"""

    st.set_page_config(
        page_title="Document Risk Analyzer",
        page_icon="📄",
        layout="wide"
    )

    st.title("📄 Document Risk Analyzer")
    st.write("Upload PDF or DOCX documents to analyze for entities and financial crime risks")

    # Sidebar
    with st.sidebar:
        st.header("About")
        st.write("""
        This tool processes documents through 4 stages:

        1. **Extract & Summarize** - Generate document summary
        2. **Extract Entities** - Identify persons & companies
        3. **Describe Entities** - Generate detailed profiles
        4. **Risk Analysis** - Flag money laundering & sanctions evasion
        """)

        st.write("---")
        st.write("**Setup:**")
        st.write("Set OPENAI_API_KEY environment variable before running")

    # Get API key
    api_key = os.getenv("OPENAI_API_KEY", "")

    if not api_key:
        st.error("⚠️ Please set OPENAI_API_KEY environment variable")
        st.code("export OPENAI_API_KEY=sk-your-key-here")
        st.stop()

    # File upload
    uploaded_files = st.file_uploader(
        "Upload Documents (PDF or DOCX)",
        type=["pdf", "docx"],
        accept_multiple_files=True
    )

    if uploaded_files:
        st.write(f"**{len(uploaded_files)} file(s) uploaded**")

        if st.button("🚀 Process Documents", type="primary"):
            all_results = []

            for uploaded_file in uploaded_files:
                # Save to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_path = tmp_file.name

                try:
                    result = process_document(api_key, tmp_path, uploaded_file.name)
                    if result:
                        all_results.append(result)
                finally:
                    # Clean up temporary file
                    os.unlink(tmp_path)

                st.write("---")

            st.success(f"✅ Processed {len(all_results)} document(s) successfully!")

            # Download all results
            if len(all_results) > 1:
                st.download_button(
                    label="📥 Download All Results (JSON)",
                    data=json.dumps(all_results, indent=2),
                    file_name="all_results.json",
                    mime="application/json"
                )


if __name__ == "__main__":
    main()
