"""Module for formatting EDD assessment and KYC quality check results into a Word report."""

import os

from constants import OUTPUT_FOLDER
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

from main.utils.logger_config import setup_logger

logger = setup_logger(__name__)


class output_writer:
    """Class for formatting EDD assessment and KYC quality check results into a Word report."""

    def create_output_folder(self):
        """Create an output folder with the current date and time."""
        self.output_folder = OUTPUT_FOLDER
        os.makedirs(self.output_folder, exist_ok=True)

    def get_display_name(self, key: str, is_req_c: bool = False, is_edd: bool = False):
        """Return the display name of the given key."""

        display_names_edd = {
            "type_of_business_relationship": "1. Type of Business Relationship",
            "request_type": "2. Request Type",
            "risk_category": "3. Risk Category",
            "purpose_of_business_relationship": "4. Purpose of Business Relationship",
            "expected_nnm_or_current_aum": (
                "5. Current AUM" if is_req_c else "5. Expected NNM"
            ),
            "expected_transactions": (
                "6. Transactions" if is_req_c else "6. Expected Transactions"
            ),
            "activity": "7. Activity",
            "total_wealth_composition": "8. Total Wealth Composition",
            "source_of_wealth": "9. Source of Wealth",
            "corroboration": "10. Corroboration",
            "negative_news": "11. Negative News and Assessment/Screening Results",
            "other_risk_aspects": "12. Other Risk Aspects and Potentially Mitigating Factors",
            "other_relevant_information": "13. Other Relevant Information",
            "final_summary": "14. Final Summary",
        }

        if is_edd:
            return display_names_edd.get(key, key.replace("_", " ").capitalize())

    def create_kyc_table(self, doc: Document, kyc_profiles_output: dict):
        """Create a single KYC table.

        All sections use the last partner's data (most complete, since kyc_checks_output
        is a shared object updated sequentially). Section 11.2 shows one subtable per partner.

        Args:
            doc: Document
            kyc_profiles_output: dict {partner_name: {check_key: {status, reason, display_name}}}
        """
        # Use last partner as the reference — it holds the most up-to-date shared checks
        reference_checks = list(kyc_profiles_output.values())[-1]
        partner_names = list(kyc_profiles_output.keys())

        num_rows = len(reference_checks)
        if num_rows == 0:
            return

        doc.add_paragraph()
        table = doc.add_table(rows=num_rows, cols=3, style="Table Grid")
        table.style.paragraph_format.space_after = Pt(5)

        original_width = 1.4 * (table.columns[0].cells[0].width)
        original_width -= Cm(0.5)
        for cell in table.columns[0].cells:
            cell.width = Cm(0.5)
        for cell in table.columns[-2].cells:
            cell.width = 0.6 * cell.width
        for cell in table.columns[-1].cells:
            cell.width += original_width

        for i, (key, value) in enumerate(reference_checks.items()):
            row_cells = table.rows[i].cells

            # Display name
            display_name = (
                value.get("display_name", key.replace("_", " ").capitalize())
                if isinstance(value, dict)
                else key.replace("_", " ").capitalize()
            )
            row_cells[1].add_paragraph().add_run(display_name).bold = True

            # Status color
            check_status = value.get("status", False) if isinstance(value, dict) else False
            colour = RGBColor(0, 255, 0) if check_status else RGBColor(255, 0, 0)
            reason_text = value.get("reason", "") if isinstance(value, dict) else str(value)
            if reason_text == "Out of scope currently.":
                colour = RGBColor(211, 211, 211)
            self.set_cell_color(row_cells[0], colour)

            # Content
            if key == "consistency_checks_within_kyc_contradiction_checks":
                # Section 11.2: one subtable per partner
                cell = row_cells[2]
                cell.add_paragraph("Please refer to the annex for LLM analysis details.\n")
                for partner_name in partner_names:
                    partner_value = kyc_profiles_output[partner_name].get(key, {})
                    raw_checks = partner_value.get("raw_checks", {}) if isinstance(partner_value, dict) else {}
                    checks = [
                        {
                            "check": k.replace("_", " ").title(),
                            "contradictions_present": str(v.get("contradictory", "No")) == "Yes",
                        }
                        for k, v in raw_checks.items()
                    ]
                    cell.add_paragraph().add_run(partner_name).bold = True
                    if checks:
                        self.add_subtable(cell, checks)
            else:
                # All other sections: use reference data only (shared across partners)
                self.write_bold_instances(row_cells[2], reason_text.strip())

    def create_table(
        self,
        doc: Document,
        processed_profile_output: dict,
        is_edd: bool = False,
    ):
        """Populate the given document with EDD assessment results."""
        print(f"create_table called: is_edd={is_edd}, num_keys={len(processed_profile_output)}")
        cols = 2
        doc.add_paragraph()
        skip_keys = {"dict_parsed_text", "raw_data", "file_path", "raw_text"}

        num_rows = sum(1 for k in processed_profile_output if k not in skip_keys)
        if num_rows == 0:
            return

        table = doc.add_table(rows=num_rows, cols=cols, style="Table Grid")
        table.style.paragraph_format.space_after = Pt(5)

        is_req_c = "Request Type C" in processed_profile_output.get("request_type", "")
        original_width = 1.4 * (table.columns[0].cells[0].width)

        for cell in table.columns[-2].cells:
            cell.width = 0.6 * cell.width
        for cell in table.columns[-1].cells:
            cell.width += original_width

        i = 0
        for key, value in processed_profile_output.items():
            if key in skip_keys:
                continue

            row_cells = table.rows[i].cells
            i += 1

            display_name = self.get_display_name(key, is_req_c, is_edd=True)
            row_cells[0].add_paragraph().add_run(display_name).bold = True

            if key not in ["activity", "total_wealth_composition"]:
                self.write_bold_instances(
                    row_cells[1], value.strip() if isinstance(value, str) else str(value)
                )

            elif key == "activity":
                cell = row_cells[1]
                for partner_name, activities in processed_profile_output.get(
                    "raw_data", {}
                ).get("activity", {}).items():
                    _ = cell.add_paragraph().add_run(f"{partner_name}\n").bold = True
                    self.add_subtable(cell, activities)

            elif key == "total_wealth_composition":
                cell = row_cells[1]
                for partner_name, twc_dict in processed_profile_output.get(
                    "raw_data", {}
                ).get("total_wealth_composition", {}).items():
                    cell.add_paragraph().add_run(f"\n{partner_name}\n").bold = True
                    sub_table_keys = ["bankable_assets", "private_equity", "real_estate", "other_assets"]
                    sub_table_names = ["Bankable assets", "Private equity assets", "Real estate assets", "Other assets"]
                    for sub_key, name in zip(sub_table_keys, sub_table_names):
                        cell.add_paragraph(name + ":")
                        self.add_subtable(cell, twc_dict.get(sub_key, {}).get("assets", []))

    def add_subtable(self, cell, values):
        if not values:
            return

        column_names = [k.title().replace("_", " ") for k in values[0].keys()]
        sub_table = cell.add_table(rows=len(values) + 1, cols=len(column_names))
        sub_table.style = "Light Grid"
        sub_table.style.font.size = Pt(10)

        for _cell, name in zip(sub_table.rows[0].cells, column_names):
            _cell.text = name

        for idx, elems in enumerate(values, start=1):
            for _cell, (k, v) in zip(sub_table.rows[idx].cells, elems.items()):
                _cell.text = str(v).title().replace("_", " ")
                if _cell.text == "True":
                    self.set_cell_color(_cell, "#DC3939")
                elif _cell.text == "False":
                    self.set_cell_color(_cell, "#29B95C")

    def set_cell_color(self, cell, color):
        """Apply a solid background colour to a table cell."""
        if isinstance(color, str):
            hex_color = color.lstrip("#")
        else:
            if hasattr(color, "rgb"):
                hex_color = f"{color.rgb:06x}"
            elif isinstance(color, (int,)):
                hex_color = f"{color:06x}"
            elif isinstance(color, (list, tuple)) and len(color) == 3:
                r, g, b = color
                hex_color = f"{r:02x}{g:02x}{b:02x}"
            else:
                raise TypeError(
                    f"Unable to convert colour {color!r} to a hex string. "
                    "Supported inputs: RGBColor, (r,g,b) tuple/list, int, or hex string."
                )

        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)

    def write_bold_instances(self, document: Document, text: str):
        """Write a given text, bolding strings surrounded by double asterisks (**like this**)."""
        for para in text.replace("\n\n", "\n").split("\n"):
            is_bold = False
            run = document.add_paragraph()
            run.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
            for chunk in para.split("**"):
                _ = run.add_run(chunk)
                _.bold = is_bold
                is_bold = not is_bold

    def write_word_doc(
        self,
        edd_profile_output: dict,
        kyc_profiles_output: dict,
        case_number: str,
        partner_mappings: dict,
        edd_case: dict,
    ):
        """Write the EDD assessment and KYC quality check results in a Word report."""
        doc = Document()

        doc.add_heading(
            "EDD Assessment/KYC checks - Business Relationship " + case_number + "\n",
            level=1,
        )

        doc.add_paragraph("EDD case parners & roles", style="Subtitle")
        doc.add_paragraph(
            edd_case["contractual_partner_information"]["name"]
            + " - contractual partner,\n"
            + ",\n".join(
                [
                    item["name"] + " - " + item["role"].replace("Has ", "")
                    for item in edd_case["role_holders_information"]
                ]
            )
        )

        if partner_mappings:
            doc.add_paragraph("KYC Partners", style="Subtitle")
            doc.add_paragraph(
                "With KYC history:\n"
                + ",\n".join(
                    [" - " + item for item in partner_mappings["partner_folder_names_with_kyc"]]
                )
            )
            doc.add_paragraph(
                "Without KYC history:\n"
                + ",\n".join(
                    [" - " + item for item in partner_mappings["partner_folder_names_without_kyc"]]
                )
            )
            doc.add_paragraph(
                "Without EDD entry:\n"
                + ",\n".join(
                    [
                        " - " + item["kyc_partner_name"]
                        for item in partner_mappings["mappings"]
                        if item["match_status"] == "unmatched"
                    ]
                )
            )

        if kyc_profiles_output:
            kyc_heading = doc.add_heading("KYC Checks", level=2)
            kyc_heading.alignment = 1
            self.create_kyc_table(doc, kyc_profiles_output)

        if edd_profile_output:
            doc.add_page_break()
            edd_heading = doc.add_heading("EDD Assessment", level=2)
            edd_heading.alignment = 1
            self.create_table(doc, edd_profile_output, is_edd=True)

        # Annex section - one entry per partner
        if kyc_profiles_output:
            doc.add_page_break()
            doc.add_heading("Annex", level=2).alignment = 1
            doc.add_paragraph().add_run("KYC Check 11.2 Analysis").bold = True

            for partner_name, partner_checks in kyc_profiles_output.items():
                logger.info(f"Processing annex for partner: {partner_name}")

                contradiction_value = partner_checks.get(
                    "consistency_checks_within_kyc_contradiction_checks", {}
                )
                raw_checks = contradiction_value.get("raw_checks", {})
                logger.debug(f"raw_checks for partner {partner_name}: {raw_checks}")

                checks = [
                    {
                        "Field": k.replace("_", " ").title(),
                        "Contradictions Present": "Yes" if v.get("contradictory", "").lower() == "true" else "No",
                        "Reasoning": v.get("reasoning", "No reasoning available."),
                    }
                    for k, v in raw_checks.items()
                ]

                partner_heading = doc.add_heading(partner_name, level=3)
                partner_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if checks:
                    logger.info(f"Creating table for partner: {partner_name}")
                    table = doc.add_table(rows=len(checks) + 1, cols=3, style="Light Grid")

                    headers = ["Field", "Contradictions Present", "Reasoning"]
                    for _cell, header in zip(table.rows[0].cells, headers):
                        _cell.text = header

                    for idx, check in enumerate(checks, start=1):
                        row_cells = table.rows[idx].cells
                        row_cells[0].text = check["Field"]
                        row_cells[1].text = check["Contradictions Present"]
                        row_cells[2].text = check["Reasoning"]
                else:
                    doc.add_paragraph(
                        f"No contradiction checks available for partner: {partner_name}."
                    )

                doc.add_paragraph()
                self.write_bold_instances(doc, contradiction_value.get("reason", ""))
                doc.add_paragraph()

        # Save the document
        save_path = os.path.join(self.output_folder, case_number)
        run_date = OUTPUT_FOLDER.split("/")[-1]
        doc.save(os.path.join(save_path, f"Report_BR_{case_number}_{run_date}.docx"))
