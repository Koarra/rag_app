
"""
Document Text Extractor with OCR Support
Supports DOCX and PDF files (with OCR for images)
"""

import os
import sys
from pathlib import Path
from typing import List, Union
import argparse

# Required libraries (install via pip):
# pip install python-docx PyPDF2 pdf2image pytesseract Pillow

try:
    from docx import Document
    import PyPDF2
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
except ImportError as e:
    print(f"Missing required library: {e}")
    print("Install with: pip install python-docx PyPDF2 pdf2image pytesseract Pillow")
    sys.exit(1)


class DocumentProcessor:
    """Process DOCX and PDF files to extract text content."""
    
    def __init__(self, use_ocr: bool = True):
        self.use_ocr = use_ocr
        
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
        
        except Exception as e:
            return f"Error processing DOCX file {file_path}: {str(e)}"
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file. Uses OCR if text extraction fails.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        text_content = []
        
        try:
            # First, try to extract text directly
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    
                    if text and text.strip():
                        text_content.append(text)
                    elif self.use_ocr:
                        # If no text or minimal text, use OCR
                        print(f"Using OCR for page {page_num + 1} of {file_path}")
                        ocr_text = self._ocr_pdf_page(file_path, page_num)
                        if ocr_text:
                            text_content.append(ocr_text)
            
            # If no text extracted and OCR is enabled, try OCR on entire PDF
            if not text_content and self.use_ocr:
                print(f"No text found in PDF. Running full OCR on {file_path}")
                text_content.append(self._ocr_entire_pdf(file_path))
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            return f"Error processing PDF file {file_path}: {str(e)}"
    
    def _ocr_pdf_page(self, pdf_path: str, page_num: int) -> str:
        """
        Perform OCR on a specific PDF page.
        
        Args:
            pdf_path: Path to PDF file
            page_num: Page number (0-indexed)
            
        Returns:
            OCR extracted text
        """
        try:
            # Convert specific page to image
            images = convert_from_path(
                pdf_path, 
                first_page=page_num + 1, 
                last_page=page_num + 1,
                dpi=300
            )
            
            if images:
                return pytesseract.image_to_string(images[0])
            return ""
        
        except Exception as e:
            print(f"OCR error on page {page_num + 1}: {str(e)}")
            return ""
    
    def _ocr_entire_pdf(self, pdf_path: str) -> str:
        """
        Perform OCR on entire PDF.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            OCR extracted text from all pages
        """
        try:
            images = convert_from_path(pdf_path, dpi=300)
            text_parts = []
            
            for i, img in enumerate(images):
                print(f"OCR processing page {i + 1}/{len(images)}")
                text = pytesseract.image_to_string(img)
                if text.strip():
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            print(f"Full OCR error: {str(e)}")
            return ""
    
    def process_file(self, file_path: str) -> str:
        """
        Process a single file and extract text.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text with file header
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return f"Error: File not found: {file_path}"
        
        ext = file_path.suffix.lower()
        
        print(f"Processing: {file_path.name}")
        
        header = f"\n{'='*80}\nFile: {file_path.name}\n{'='*80}\n\n"
        
        if ext == '.docx':
            content = self.extract_from_docx(str(file_path))
        elif ext == '.pdf':
            content = self.extract_from_pdf(str(file_path))
        else:
            content = f"Unsupported file format: {ext}"
        
        return header + content
    
    def process_files(self, file_paths: List[str]) -> str:
        """
        Process multiple files and combine their text.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            Combined text from all files
        """
        results = []
        
        for file_path in file_paths:
            results.append(self.process_file(file_path))
        
        return "\n\n".join(results)


def main():
    """Main function with CLI interface."""
    parser = argparse.ArgumentParser(
        description='Extract text from DOCX and PDF files with OCR support'
    )
    parser.add_argument(
        'files',
        nargs='+',
        help='One or more files to process (DOCX or PDF)'
    )
    parser.add_argument(
        '--no-ocr',
        action='store_true',
        help='Disable OCR for PDFs'
    )
    parser.add_argument(
        '-o', '--output',
        help='Output file path (default: print to stdout)'
    )
    
    args = parser.parse_args()
    
    # Initialize processor
    processor = DocumentProcessor(use_ocr=not args.no_ocr)
    
    # Process files
    extracted_text = processor.process_files(args.files)
    
    # Output results
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        print(f"\nText extracted and saved to: {args.output}")
    else:
        print(extracted_text)


if __name__ == "__main__":
    # Example usage in code:
    # processor = DocumentProcessor(use_ocr=True)
    # text = processor.process_files(['document.pdf', 'report.docx'])
    # print(text)
    
    main()
